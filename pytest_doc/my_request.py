import requests
import docx
import json
from docx.shared import RGBColor
import os
import my_doc


#接口配置
def auto_request(url,headers,my_token,doc,ymd,hms):
    nothing_data = [
        "/xearth/epidemic-daily/list",
        "/xearth/risk-type/click-type",
        "/xearth/project/project-statistics",
        "/xearth/plague/project-v2",
        "/xearth/user/index",
        "/xearth/project-file/get-file-menu",
        "/xearth/statistics/sentiment-tree-v2",
        "/xearth/country/country-v2",
        "/xearth/report/weekly-index-v2",
        "/xearth/statistics/onekey-alarm",
        "/xearth/statistics/receive-police-list-one",
        "/xearth/emergency/process-list"]
    a = len(nothing_data)
    #生成接口返回码的列表
    lists = []
    table = doc.add_table(rows=1, cols=2, style = "Light Shading Accent 3")
    my_table = table.rows[0].cells
    my_table[0].text = "接口名称"
    my_table[1].text = "返回code码"
    for i in range(a):
        no_data_json = requests.request("POST", url + nothing_data[i] + "?access-token=" + my_token, headers=headers)
        lists.append(nothing_data[i])
        lists.append(no_data_json.status_code)
        #生成详细接口数据
        test = doc.add_paragraph(style='List Number').add_run(no_data_json.text)
        if (no_data_json.status_code != 200):
            test.font.color.rgb = RGBColor(250, 0, 0)
        doc.save('D:\\'+ymd+hms+'接口监控.docx')

    have_data = (
        "/xearth/risk-type/get-type",
        "/xearth/area-list-new/get-area-risk-v2"
    )
    have_data_data = [
        '"{\"source\":\"all\"}"',
        '"{\"area_id\":0,\"type_id\":\"17\"}"'
    ]
    b = len(have_data)
    for i in range(b):
        have_data_json = requests.request("POST", url + have_data[i] + "?access-token=" + my_token,data=have_data_data[i],headers=headers)
        lists.append(have_data[i])
        lists.append(have_data_json.status_code)
        test = doc.add_paragraph(style='List Number').add_run(have_data_json.text)
        if(have_data_json.status_code != 200):
            test.font.color.rgb = RGBColor(250,0,0)
        doc.save('D:\\'+ymd+hms+'接口监控.docx')
    #将返回数据写入到列表
    tup_a = tuple(lists)
    lst = []
    for i in range(0, len(tup_a), 2):
        lst.append((tup_a[i:i + 2]))
    tup_b = tuple(lst)
    for a, b in tup_b:
        abc = table.add_row().cells
        abc[0].text = str(a)
        abc[1].text = str(b)
    doc.save('D:\\' + ymd + hms + '接口监控.docx')