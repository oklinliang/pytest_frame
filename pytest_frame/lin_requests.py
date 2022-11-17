# -*- coding:utf-8 -*-
import lin_token
import requests
import time
import lin_main
import lin_config
import docx
from docx.shared import RGBColor

#接口配置
def auto_request():
    all_list = []
    err_list = []
    doc = docx.Document()
    #无参数请求接口
    nothing_data = [

        ]
    nothing_data_len = len(nothing_data)
    table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 3")
    my_table = table.rows[0].cells
    my_table[0].text = "接口名称"
    my_table[1].text = "返回code码"
    for i in range(nothing_data_len):
        no_data_json = requests.request("POST", lin_config.request_url + nothing_data[i] + "?access-token=" + lin_token.login_token(), headers=lin_config.request_headers)
        all_list.append(nothing_data[i])
        all_list.append(no_data_json.status_code)
        # 生成详细接口数据
        test = doc.add_paragraph(style='List Number').add_run(no_data_json.text)
        time.sleep(2)
        if (no_data_json.status_code != 200):
            #在报告中标红
            test.font.color.rgb = RGBColor(250, 0, 0)
            #统计报错接口数
            response_time = str(no_data_json.elapsed.total_seconds())
            request_json = str(no_data_json.json())
            request_url = str(no_data_json.url)
            err_list.append("接口名称："+ request_url)
            err_list.append("响应时间：" + response_time)
            err_list.append("接口返回："+ request_json)
        else:
            response_time = str(no_data_json.elapsed.total_seconds())
            request_json = str(no_data_json.json())
            request_url = str(no_data_json.url)
            response_time_float = float(response_time)
            if(response_time_float >= 2):
                err_list.append("接口名称：" + request_url)
                err_list.append("响应时间：" + response_time)
                err_list.append("接口返回：" + request_json)
    #有参数请求接口
    have_data = [

    ]
    #接口参数
    have_data_data = [

    ]
    have_data_len = len(have_data)
    for i in range(have_data_len):
        have_data_json = requests.request("POST", lin_config.request_url + have_data[i] + "?access-token=" + lin_token.login_token(),data=have_data_data[i],headers=lin_config.request_headers)
        all_list.append(have_data[i])
        all_list.append(have_data_json.status_code)
        test = doc.add_paragraph(style='List Number').add_run(have_data_json.text)
        time.sleep(2)
        if (have_data_json.status_code != 200):
            # 在报告中标红
            test.font.color.rgb = RGBColor(250, 0, 0)
            # 统计报错接口数
            response_time = str(have_data_json.elapsed.total_seconds())
            request_json = str(have_data_json.json())
            request_url = str(have_data_json.url)
            err_list.append("接口名称：" + request_url)
            err_list.append("响应时间：" + response_time)
            err_list.append("接口返回：" + request_json)
        else:
            response_time = str(have_data_json.elapsed.total_seconds())
            request_json = str(have_data_json.json())
            request_url = str(have_data_json.url)
            response_time_float = float(response_time)
            if (response_time_float >= 2):
                err_list.append("接口名称：" + request_url)
                err_list.append("响应时间：" + response_time)
                err_list.append("接口返回：" + request_json)
    # 将返回数据写入到列表
    tup_a = tuple(all_list)
    lst = []
    for i in range(0, len(tup_a), 2):
        lst.append((tup_a[i:i + 2]))
    tup_b = tuple(lst)
    for a, b in tup_b:
        abc = table.add_row().cells
        abc[0].text = str(a)
        abc[1].text = str(b)
    doc.save('D:\\' + lin_main.time_ymd + lin_main.time_hms + '接口监控.docx')
    err_list_len = len(err_list)
    type(err_list)
    return err_list_len


